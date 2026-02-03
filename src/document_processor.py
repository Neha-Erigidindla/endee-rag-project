"""
Document Processor
Handles document loading, chunking, and embedding generation
"""

import os
from typing import List, Dict, Optional, Tuple
import logging
from pathlib import Path
from dataclasses import dataclass
import hashlib

# Document loaders
import PyPDF2
from docx import Document as DocxDocument

# Embedding model
from sentence_transformers import SentenceTransformer
import numpy as np

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of a document"""
    id: str
    text: str
    metadata: Dict
    embedding: Optional[List[float]] = None


class DocumentProcessor:
    """Process documents for RAG system"""
    
    def __init__(
        self,
        model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
        chunk_size: int = 512,
        chunk_overlap: int = 50
    ):
        """
        Initialize document processor
        
        Args:
            model_name: Sentence transformer model name
            chunk_size: Maximum characters per chunk
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        logger.info(f"Loading embedding model: {model_name}")
        self.model = SentenceTransformer(model_name)
        self.embedding_dim = self.model.get_sentence_embedding_dimension()
        
        logger.info(f"Model loaded. Embedding dimension: {self.embedding_dim}")
    
    def load_pdf(self, file_path: str) -> str:
        """
        Load text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                
                logger.info(f"Loaded PDF: {file_path} ({len(pdf_reader.pages)} pages)")
                return text
                
        except Exception as e:
            logger.error(f"Failed to load PDF {file_path}: {e}")
            raise
    
    def load_docx(self, file_path: str) -> str:
        """
        Load text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            logger.info(f"Loaded DOCX: {file_path} ({len(doc.paragraphs)} paragraphs)")
            return text
            
        except Exception as e:
            logger.error(f"Failed to load DOCX {file_path}: {e}")
            raise
    
    def load_txt(self, file_path: str) -> str:
        """
        Load text from TXT file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            File contents
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"Loaded TXT: {file_path}")
            return text
            
        except Exception as e:
            logger.error(f"Failed to load TXT {file_path}: {e}")
            raise
    
    def load_document(self, file_path: str) -> str:
        """
        Load document based on file extension
        
        Args:
            file_path: Path to document
            
        Returns:
            Extracted text
        """
        ext = Path(file_path).suffix.lower()
        
        loaders = {
            '.pdf': self.load_pdf,
            '.docx': self.load_docx,
            '.txt': self.load_txt,
            '.md': self.load_txt
        }
        
        if ext not in loaders:
            raise ValueError(f"Unsupported file type: {ext}")
        
        return loaders[ext](file_path)
    
    def chunk_text(self, text: str, metadata: Optional[Dict] = None) -> List[str]:
        """
        Split text into overlapping chunks
        
        Args:
            text: Text to chunk
            metadata: Optional metadata to attach
            
        Returns:
            List of text chunks
        """
        if not text.strip():
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            # Get chunk
            end = start + self.chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence boundary
            if end < text_length:
                # Look for sentence ending
                for delimiter in ['. ', '.\n', '! ', '? ']:
                    last_delim = chunk.rfind(delimiter)
                    if last_delim > self.chunk_size * 0.5:  # At least 50% of chunk
                        chunk = chunk[:last_delim + 1]
                        break
            
            chunks.append(chunk.strip())
            
            # Move start position with overlap
            start = start + self.chunk_size - self.chunk_overlap
        
        logger.debug(f"Created {len(chunks)} chunks from text")
        return chunks
    
    def generate_chunk_id(self, text: str, source: str, index: int) -> str:
        """
        Generate unique ID for chunk
        
        Args:
            text: Chunk text
            source: Source document
            index: Chunk index
            
        Returns:
            Unique chunk ID
        """
        # Create hash from text for uniqueness
        text_hash = hashlib.md5(text.encode()).hexdigest()[:8]
        source_name = Path(source).stem
        
        return f"{source_name}_chunk{index}_{text_hash}"
    
    def process_document(
        self,
        file_path: str,
        additional_metadata: Optional[Dict] = None
    ) -> List[DocumentChunk]:
        """
        Process entire document into chunks with embeddings
        
        Args:
            file_path: Path to document
            additional_metadata: Extra metadata to attach
            
        Returns:
            List of document chunks with embeddings
        """
        logger.info(f"Processing document: {file_path}")
        
        # Load document
        text = self.load_document(file_path)
        
        # Chunk text
        chunks_text = self.chunk_text(text)
        
        # Generate embeddings
        logger.info(f"Generating embeddings for {len(chunks_text)} chunks...")
        embeddings = self.model.encode(
            chunks_text,
            show_progress_bar=True,
            convert_to_numpy=True
        )
        
        # Create chunk objects
        chunks = []
        for idx, (chunk_text, embedding) in enumerate(zip(chunks_text, embeddings)):
            # Base metadata
            metadata = {
                "source": os.path.basename(file_path),
                "source_path": file_path,
                "chunk_index": idx,
                "total_chunks": len(chunks_text),
                "text": chunk_text,
                "char_count": len(chunk_text)
            }
            
            # Add additional metadata
            if additional_metadata:
                metadata.update(additional_metadata)
            
            # Generate unique ID
            chunk_id = self.generate_chunk_id(chunk_text, file_path, idx)
            
            chunks.append(DocumentChunk(
                id=chunk_id,
                text=chunk_text,
                metadata=metadata,
                embedding=embedding.tolist()
            ))
        
        logger.info(f"✓ Processed {file_path}: {len(chunks)} chunks created")
        return chunks
    
    def embed_query(self, query: str) -> List[float]:
        """
        Generate embedding for search query
        
        Args:
            query: Search query text
            
        Returns:
            Query embedding vector
        """
        embedding = self.model.encode(query, convert_to_numpy=True)
        return embedding.tolist()
    
    def batch_process_directory(
        self,
        directory: str,
        file_pattern: str = "*.*"
    ) -> List[DocumentChunk]:
        """
        Process all documents in a directory
        
        Args:
            directory: Directory path
            file_pattern: Glob pattern for files
            
        Returns:
            List of all document chunks
        """
        all_chunks = []
        directory_path = Path(directory)
        
        # Find all matching files
        files = list(directory_path.glob(file_pattern))
        supported_exts = {'.pdf', '.docx', '.txt', '.md'}
        files = [f for f in files if f.suffix.lower() in supported_exts]
        
        logger.info(f"Found {len(files)} documents to process")
        
        for file_path in files:
            try:
                chunks = self.process_document(str(file_path))
                all_chunks.extend(chunks)
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                continue
        
        logger.info(f"✓ Processed {len(files)} files: {len(all_chunks)} total chunks")
        return all_chunks


# Example usage
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    
    # Initialize processor
    processor = DocumentProcessor()
    
    # Process a single document
    chunks = processor.process_document("sample.pdf")
    
    print(f"Created {len(chunks)} chunks")
    print(f"Embedding dimension: {len(chunks[0].embedding)}")
    
    # Print first chunk
    print(f"\nFirst chunk:")
    print(f"ID: {chunks[0].id}")
    print(f"Text: {chunks[0].text[:200]}...")
    print(f"Metadata: {chunks[0].metadata}")
